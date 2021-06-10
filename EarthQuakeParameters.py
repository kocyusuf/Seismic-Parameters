import EQParameters as eqp
import turkeyMap as maptr
import drawSpectralGraphs as drawSpect
import sys
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtWidgets import QFileDialog
import webbrowser
import pandas as pd
import math
from docx import Document
from docx.shared import Inches
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys

def round_down(n, decimals=0):
    multiplier = 10 ** decimals
    return math.floor(n * multiplier) / multiplier
def nearest_number(n, querryList):
    myList = querryList
    nNum = min(myList, key=lambda x:abs(x-n))
    myList.remove(nNum)
    nNum2 = min(myList, key=lambda x:abs(x-n))
    return nNum, nNum2
class EarthQuakeMap(QtWidgets.QMainWindow):
    def __init__(self):
        super(EarthQuakeMap, self).__init__()
        self.uimenu = eqp.Ui_Form()
        self.uimenu.setupUi(self)
        self.EQLevels = ['DD1', 'DD2', 'DD3', 'DD4']
        self.uimenu.cmboxEQLevels.addItems(self.EQLevels)
        self.uimenu.cmboxGroundClass.addItems(['ZA', 'ZB', 'ZC', 'ZD', 'ZE'])
        self.uimenu.btnGetCoordOnMap.clicked.connect(self.getCoordOnMap)
        self.uimenu.btnCalculate.clicked.connect(self.calculateParameters)
        self.uimenu.btnCreateReport.clicked.connect(self.createSeismicReport)
    def saveReport(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getSaveFileName(self, "Kaydet", "",
                                                  "All Files (*);;Word Files (*.docx)", options=options)
        self.show()
        self.document.save(f'{fileName}.docx')
    def getCoordOnMap(self):
        maptr.indexMap()
        webbrowser.open('trindex.html')
    def calculateParameters(self):
        """ Get User Input Values """
        Latitute = float(self.uimenu.txtLatitute.text())
        Longitude = float(self.uimenu.txtLongitude.text())
        driver_path = "chromedriver.exe"
        browser = webdriver.Chrome(executable_path=driver_path)
        browser.get('https://www.google.com/maps')
        delay = 5
        inputElement = browser.find_element_by_id("searchboxinput")
        inputElement.send_keys(f'{Latitute, Longitude}')
        inputElement.send_keys(Keys.ENTER)
        inputElement.submit()
        time.sleep(delay)
        browser.save_screenshot('location.png')
        browser.quit()
        """ Round Longitude Value """
        LongCeil = round_down(Longitude, 1)
        LongValue = round(LongCeil + 0.05, 2)
        print("long", LongValue)
        """ Round Latitute Value """
        LatiRoundValue = round(Latitute, 3)
        LatiFloor = round(round(Latitute, 1) - 0.05, 2)
        print('LatiFloor: ', LatiFloor)
        LatiCeil = round(round(Latitute, 1) + 0.05, 2)
        print("LatCeil", LatiCeil)
        """ Create DataFrames """
        getPGA = pd.read_excel("paramPga.xlsx", engine='openpyxl')
        getPGV = pd.read_excel('paramPgv.xlsx', engine='openpyxl')
        getS1 = pd.read_excel('paramS1.xlsx', engine='openpyxl')
        getSs = pd.read_excel('paramSs.xlsx', engine='openpyxl')
        print("Veri setleri oluşturuldu.")
        # getParamPGA = getPGA.loc[(getPGA['Boylam'] == Longitude) & (getPGA['Enlem'] == Latitute), ['DD2']]
        # getParamPGA = getPGA.loc[(getPGA['Boylam'] == LongValue) & (getPGA['Enlem'] == LatiValue)]
        """ Get PGA Values """
        getParamPGACeilCon = (getPGA['Boylam'] == LongValue) & (getPGA['Enlem'] == LatiCeil)
        getParamPGAFloorCon = (getPGA['Boylam'] == LongValue) & (getPGA['Enlem'] == LatiFloor)
        indexPGACeil = list(getPGA.loc[getParamPGACeilCon].index)
        indexPGAFloor = list(getPGA.loc[getParamPGAFloorCon].index)
        print('index:', indexPGAFloor)
        pgaCeilDict = {}
        for level in self.EQLevels:
            pga = getPGA._get_value(indexPGACeil[0], level)
            pgaCeilDict[level] = pga
        print('pgaceilDict:', pgaCeilDict)
        pgaFloorDict = {}
        for level in self.EQLevels:
            pga = getPGA.loc[indexPGAFloor[0], level]
            pgaFloorDict[level] = pga
        print('pgafloorDict:', pgaFloorDict)
        self.valuesPGADict = {}
        for level in self.EQLevels:
            pgaValue = pgaCeilDict[level] + (((pgaFloorDict[level] - pgaCeilDict[level]) * (LatiRoundValue - LatiCeil)) / (LatiFloor - LatiCeil))
            self.valuesPGADict[level] = pgaValue
        print("well done pga")
        """ Get PGV Values """
        getParamPGVCeilCon = (getPGV['Boylam'] == LongValue) & (getPGV['Enlem'] == LatiCeil)
        getParamPGVFloorCon = (getPGV['Boylam'] == LongValue) & (getPGV['Enlem'] == LatiFloor)
        indexPGVCeil = list(getPGV.loc[getParamPGVCeilCon].index)
        indexPGVFloor = list(getPGV.loc[getParamPGVFloorCon].index)
        pgvCeilDict = {}
        for level in self.EQLevels:
            pgv = getPGV._get_value(indexPGVCeil[0], level)
            pgvCeilDict[level] = pgv
        pgvFloorDict = {}
        for level in self.EQLevels:
            pgv = getPGV.loc[indexPGVFloor[0], level]
            pgvFloorDict[level] = pgv
        self.valuesPGVDict = {}
        for level in self.EQLevels:
            pgvValue = pgvCeilDict[level] + (
                        ((pgvFloorDict[level] - pgvCeilDict[level]) * (LatiRoundValue - LatiCeil)) / (
                            LatiFloor - LatiCeil))
            self.valuesPGVDict[level] = pgvValue
        print("well done pgv")
        """ Get S1 Values """
        getParamS1CeilCon = (getS1['Boylam'] == LongValue) & (getS1['Enlem'] == LatiCeil)
        getParamS1FloorCon = (getS1['Boylam'] == LongValue) & (getS1['Enlem'] == LatiFloor)
        indexS1Ceil = list(getS1.loc[getParamS1CeilCon].index)
        indexS1Floor = list(getS1.loc[getParamS1FloorCon].index)
        s1CeilDict = {}
        for level in self.EQLevels:
            s1 = getS1._get_value(indexS1Ceil[0], level)
            s1CeilDict[level] = s1
        s1FloorDict = {}
        for level in self.EQLevels:
            s1 = getS1.loc[indexS1Floor[0], level]
            s1FloorDict[level] = s1
        self.valuesS1Dict = {}
        for level in self.EQLevels:
            s1Value = s1CeilDict[level] + (
                        ((s1FloorDict[level] - s1CeilDict[level]) * (LatiRoundValue - LatiCeil)) / (
                            LatiFloor - LatiCeil))
            self.valuesS1Dict[level] = s1Value
        print("well done s1")
        """ Get Ss Values """
        getParamSsCeilCon = (getSs['Boylam'] == LongValue) & (getSs['Enlem'] == LatiCeil)
        getParamSsFloorCon = (getSs['Boylam'] == LongValue) & (getSs['Enlem'] == LatiFloor)
        indexSsCeil = list(getSs.loc[getParamSsCeilCon].index)
        indexSsFloor = list(getSs.loc[getParamSsFloorCon].index)
        ssCeilDict = {}
        for level in self.EQLevels:
            ss = getSs._get_value(indexSsCeil[0], level)
            ssCeilDict[level] = ss
        ssFloorDict = {}
        for level in self.EQLevels:
            ss = getSs.loc[indexSsFloor[0], level]
            ssFloorDict[level] = ss
        self.valuesSsDict = {}
        for level in self.EQLevels:
            ssValue = ssCeilDict[level] + (
                        ((ssFloorDict[level] - ssCeilDict[level]) * (LatiRoundValue - LatiCeil)) / (
                            LatiFloor - LatiCeil))
            self.valuesSsDict[level] = ssValue
        print("well done Ss")
        print("PGA :", self.valuesPGADict)
        print("PGV: ", self.valuesPGVDict)
        print("S1 :", self.valuesS1Dict)
        print("SS :", self.valuesSsDict)
        self.uimenu.linePGA.setText(f"{str(round(self.valuesPGADict[self.uimenu.cmboxEQLevels.currentText()], 3))}")
        self.uimenu.linePGV.setText(f"{str(round(self.valuesPGVDict[self.uimenu.cmboxEQLevels.currentText()], 3))}")
        self.uimenu.lineS1.setText(f"{str(round(self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()], 3))}")
        self.uimenu.lineSs.setText(f"{str(round(self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()], 3))}")
        """ Calculate Design Spectral Acceleration Coefficents """
        getF1 = pd.read_excel('paramF1.xlsx')
        getFs = pd.read_excel('paramFs.xlsx')
        S1QuerryList = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6]
        SsQuerryList = [0.25, 0.50, 0.75, 1.0, 1.25, 1.5]
        """ Get F1 Value """
        getParamF1 = (getF1['Zemin'] == self.uimenu.cmboxGroundClass.currentText())
        indexF1 = list(getF1.loc[getParamF1].index)
        if self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()] <= 0.10:
            self.valueF1 = getF1._get_value(indexF1[0], '0.1')
        elif self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()] >= 0.60:
            self.valueF1 = getF1._get_value(indexF1[0], '0.6')
        else:
            f1Dict = {}
            for value in S1QuerryList:
                f1 = getF1._get_value(indexF1[0], value)
                f1Dict[value] = f1
            nNum, nNum2 = nearest_number(self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()], S1QuerryList)
            for v in S1QuerryList:
                self.valueF1 = f1Dict[nNum2] + ((self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()] - nNum2) * (
                            f1Dict[nNum] - f1Dict[nNum2]) / (nNum - nNum2))
        print("F1: ", self.valueF1)
        """ Get Fs Value """
        getParamFs = (getFs['Zemin'] == self.uimenu.cmboxGroundClass.currentText())
        indexFs = list(getFs.loc[getParamFs].index)
        if self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()] <= 0.25:
            self.valueFs = getFs._get_value(indexFs[0], '0.25')
        elif self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()] >= 1.50:
            self.valueFs = getFs._get_value(indexFs[0], '1.5')
        else:
            fsDict = {}
            for value in SsQuerryList:
                fs = getFs._get_value(indexFs[0], value)
                fsDict[value] = fs
            nNums, nNum2s = nearest_number(self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()], SsQuerryList)
            for v in SsQuerryList:
                self.valueFs = fsDict[nNum2s] + ((self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()] - nNum2s) * (
                        fsDict[nNums] - fsDict[nNum2s]) / (nNums - nNum2s))
        print("Fs: ", self.valueFs)
        self.Sd1 = self.valueF1 * self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()]
        self.Sds = self.valueFs * self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()]
        print("Sd1 :", self.Sd1)
        print("Sds :", self.Sds)
        self.uimenu.lineSd1.setText(str(round(self.Sd1, 3)))
        self.uimenu.lineSds.setText(str(round(self.Sds, 3)))
        self.drawGraphics()
    def drawGraphics(self):
        drawSpect.drawHorGraphs(self.Sd1, self.Sds)
        drawSpect.drawHVerGraphs(self.Sd1, self.Sds)
        self.TB = self.Sd1 / self.Sds
        self.TA = 0.2 * self.TB
        self.TAD = self.TA / 3
        self.TBD = self.TB / 3
        self.uimenu.lblHorSpectGraph.setPixmap(QtGui.QPixmap("horiSpect.png"))
        self.uimenu.lblVerSpectGraph.setPixmap(QtGui.QPixmap("vertSpect.png"))
    def createSeismicReport(self):
        print("Rapor oluşturuluyor.")
        self.document = Document()
        self.document.add_picture('afadTDTH.png', width=Inches(5.0625), height=Inches(2.65625))
        self.document.add_heading('Sismik Parametreler Detay Raporu', 0)
        self.document.add_heading('Kullanıcı Girdileri', level=1)
        p = self.document.add_paragraph(f'Deprem yer hareketi seviyesi     ')
        p.add_run(f'{self.uimenu.cmboxEQLevels.currentText()}').bold = True
        p2 = self.document.add_paragraph(f'Yerel Zemin Sınıfı     ')
        p2.add_run(f'{self.uimenu.cmboxGroundClass.currentText()}').bold = True
        p3 = self.document.add_paragraph(f'Enlem     ')
        p3.add_run(f'{self.uimenu.txtLatitute.text()}').bold = True
        p4 = self.document.add_paragraph(f'Boylam     ')
        p4.add_run(f'{self.uimenu.txtLongitude.text()}').bold = True
        self.document.add_picture('location.png', width=Inches(6.5), height=Inches(3.5))
        self.document.add_heading('Çıktılar', level=1)
        self.document.add_paragraph(
            f'Ss = {round(self.valuesSsDict[self.uimenu.cmboxEQLevels.currentText()], 3)}\tS1 = {round(self.valuesS1Dict[self.uimenu.cmboxEQLevels.currentText()], 3)}\tPGA = {round(self.valuesPGADict[self.uimenu.cmboxEQLevels.currentText()], 3)}\tPGV = {round(self.valuesPGVDict[self.uimenu.cmboxEQLevels.currentText()], 3)}')
        self.document.add_paragraph(f'Fs = {round(self.valueFs, 3)}\tF1 = {round(self.valueF1, 3)}')
        self.document.add_paragraph('Sds = SsFs')
        self.document.add_paragraph('Sd1 = S1F1')
        self.document.add_paragraph(f'Sds = {round(self.Sds, 4)}\tSd1 = {round(self.Sd1, 4)}')
        self.document.add_paragraph(
            'Ss : Kısa periyot harita spektral ivme katsayısı[boyutsuz].\nS1 : 1 s periyot için harita spektral ivme katsayısı[boyutsuz].\nPGA : En büyük yer ivmesi [g].\nPGV : En büyük yer hızı [cm/s].\nSds : Kısa periyot tasarım spektral ivme katsayısı [boyutsuz].\nSd1 : 1 s periyot için tasarım spektral ivme katsayısı [boyutsuz].\n')
        self.document.add_heading('Yatay Elastik Tasarım Spektrumu', level=1)
        self.document.add_picture('horiSpect.png')
        self.document.add_paragraph('Sae(T) = (0.4 + 0.6T/TA)Sds          0 < T < TA')
        self.document.add_paragraph('Sae(T) = Sds                         TA < T < TB')
        self.document.add_paragraph('Sae(T) = Sd1/T                       TB < T < TL')
        self.document.add_paragraph('Sae(T) = Sd1(TL/T2)                  TL < T')
        self.document.add_paragraph('TA = 0.2(Sd1/Sds)\t\tTB = Sd1/Sds\t\tTL = 6s\n')
        self.document.add_paragraph(f'TA={round(self.TA, 3)} (s)\tTB={round(self.TB, 3)} (s)\tTL=6.000 (s)')
        self.document.add_heading('Düşey Elastik Tasarım Spektrumu', level=1)
        self.document.add_picture('vertSpect.png')
        self.document.add_paragraph('SaeD(T) = (0.32 + 0.48T/TAD)Sds          0 < T < TAD')
        self.document.add_paragraph('SaeD(T) = 0.8Sds                         TAD < T < TBD')
        self.document.add_paragraph('SaeD(T) = 0.8Sds(TBD/T)                  TBD < T < TLD')
        self.document.add_paragraph('TAD = TA/3\t\tTBD = TB/3\t\tTLD = TL/2\n')
        self.document.add_paragraph(f'TAD={round(self.TAD, 3)} (s)\tTBD={round(self.TBD, 3)} (s)\tTLD=3.000 (s)')
        self.document.add_page_break()
        self.saveReport()
def app():
    app_win = QtWidgets.QApplication(sys.argv)
    window = EarthQuakeMap()
    window.show()
    sys.exit(app_win.exec_())
app()
